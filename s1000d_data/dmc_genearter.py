import os
import json
import logging
import shutil
from datetime import datetime
from pathlib import Path
import requests
import re

# Required libraries - install with: pip install requests beautifulsoup4 python-docx lxml
try:
    from bs4 import BeautifulSoup
    import docx
except ImportError as e:
    print(f"Error: A required library is missing: {e.name}")
    print("Please install all required libraries by running:")
    print("pip install requests beautifulsoup4 python-docx lxml")
    exit()

# --- CONFIGURATION ---
OLLAMA_API_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "llama3.2:latest"
LLM_REQUEST_TIMEOUT = 300

# --- DIRECTORIES ---
INPUT_DOCS_DIR = Path("documents_to_process")
PROCESSED_DIR = INPUT_DOCS_DIR / "processed"
DATA_DIR = Path("s1000d_data")
OUTPUT_DIR = Path("output")
LOGS_DIR = Path("logs")

# --- S1000D DMC CONFIGURATION ---
DMC_MODEL_IDENT_CODE = "P15"
DMC_SYSTEM_DIFF_CODE = "A"
DMC_ASSY_CODE = "0000"
DMC_ITEM_LOCATION_CODE = "D"
DMC_DEFAULT_SYS_CODE = "00"
DMC_DEFAULT_SUB_SYS_CODE = "0"
DMC_DEFAULT_INFO_CODE = "000"

def setup_environment():
    """Creates necessary directories and dummy files."""
    for dir_path in [INPUT_DOCS_DIR, PROCESSED_DIR, DATA_DIR, OUTPUT_DIR, LOGS_DIR]:
        dir_path.mkdir(exist_ok=True)
    # dummy_files = {
    #     "system_composition.docx": ("System Composition", "This document provides an equipment list of the Armaments system."),
    #     "launcher_unpacking_guide.docx": ("Unpacking the Launcher System", "This document contains the procedure to unpack the primary launcher assembly, which is part of the guided missile system.")
    # }
    # for filename, (heading, content) in dummy_files.items():
    #     filepath = INPUT_DOCS_DIR / filename
    #     if not filepath.exists():
    #         print(f"Creating dummy document: {filepath}")
    #         doc = docx.Document()
    #         doc.add_heading(heading, level=1)
    #         doc.add_paragraph(content)
    #         doc.save(filepath)

def load_data_file(filepath, loader_func, file_type):
    """Generic function to load and parse data files."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return loader_func(f)
    except FileNotFoundError:
        logging.critical(f"{file_type} file not found at {filepath}. Aborting.")
        return None
    except Exception as e:
        logging.error(f"Error loading {file_type} from {filepath}: {e}")
        return None

def parse_sns_from_file_object(file_object):
    """Parses S1000D SNS XML, including subsystems."""
    soup = BeautifulSoup(file_object, 'xml')
    sns_data = {}
    for system in soup.find_all('snsSystem'):
        sys_code_tag, sys_title_tag = system.find('snsCode', recursive=False), system.find('snsTitle', recursive=False)
        if not (sys_code_tag and sys_title_tag and sys_code_tag.text.strip()): continue
        sys_code, sys_title = sys_code_tag.text.strip(), sys_title_tag.text.strip()
        subsystems = {}
        for subsystem in system.find_all('snsSubSystem'):
            sub_code_tag, sub_title_tag = subsystem.find('snsCode'), subsystem.find('snsTitle')
            if sub_code_tag and sub_title_tag and sub_code_tag.text.strip():
                subsystems[sub_code_tag.text.strip()] = sub_title_tag.text.strip()
        sns_data[sys_code] = {'title': sys_title, 'subsystems': subsystems}
    return sns_data

def extract_docx_structure(filepath):
    """Extracts text content from a .docx file."""
    try:
        doc = docx.Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        logging.error(f"Could not read docx file {filepath}: {e}")
        return None

def get_descriptions_from_llm(full_text):
    """(IMPROVED PROMPT) Asks LLM for specific English descriptions."""
    prompt = f"""
    You are an expert technical analyst. Your task is to read the following document content and identify three key pieces of information.
    Your response MUST be ONLY a single, flat JSON object.

    DOCUMENT CONTENT:
    "{full_text[:2000]}"

    INSTRUCTIONS:
    1.  `system_name`: Identify the main technical system being discussed (e.g., "Guided missile systems", "Structure", "Propulsion").
    2.  `subsystem_name`: From the document, identify the MOST SPECIFIC component or part being discussed (e.g., "Launchers", "Hull", "Propulsion diesel"). AVOID using "General" unless no specific component is mentioned.
    3.  `purpose_description`: Describe the document's primary purpose in a short phrase (e.g., "Procedure to unpack items", "Equipment list", "System description").

    Provide your final answer as a JSON object.
    """
    try:
        logging.info("Querying LLM for document analysis...")
        payload = {"model": OLLAMA_MODEL, "prompt": prompt, "stream": False, "format": "json"}
        response = requests.post(OLLAMA_API_URL, json=payload, timeout=LLM_REQUEST_TIMEOUT)
        response.raise_for_status()
        descriptions = json.loads(response.json().get('response', '{}'))
        logging.info(f"LLM analysis returned: {descriptions}")
        return descriptions
    except Exception as e:
        logging.error(f"LLM analysis failed: {e}")
        return None

def get_words(text):
    """Helper function to normalize text into a set of words."""
    return set(re.findall(r'\b\w+\b', text.lower()))

def find_codes_from_descriptions(descriptions, sns_data, info_codes):
    """(IMPROVED LOGIC) Uses word-set scoring for robust matching."""
    found_codes = {
        "systemCode": DMC_DEFAULT_SYS_CODE,
        "subSystemCode": DMC_DEFAULT_SUB_SYS_CODE,
        "infoCode": DMC_DEFAULT_INFO_CODE
    }
    
    # --- Find SNS Code using word-set scoring ---
    system_words = get_words(descriptions.get('system_name', ''))
    subsystem_words = get_words(descriptions.get('subsystem_name', ''))
    
    best_sys_match = (None, 0)
    for sys_code, data in sns_data.items():
        title_words = get_words(data['title'])
        score = len(system_words.intersection(title_words))
        if score > best_sys_match[1]:
            best_sys_match = (sys_code, score)
    
    if best_sys_match[0] and best_sys_match[1] > 0:
        matched_sys_code = best_sys_match[0]
        found_codes["systemCode"] = matched_sys_code
        
        best_sub_match = (None, -1) # Use -1 to allow 0-score matches for "General"
        for sub_code, sub_title in sns_data[matched_sys_code]['subsystems'].items():
            sub_title_words = get_words(sub_title)
            score = len(subsystem_words.intersection(sub_title_words))
            # Heavily prioritize "General" if the LLM suggests it.
            if "general" in subsystem_words and "general" in sub_title_words:
                score = 99
            if score > best_sub_match[1]:
                best_sub_match = (sub_code, score)
        if best_sub_match[0]:
            found_codes["subSystemCode"] = best_sub_match[0]

    # --- Find Info Code using word-set scoring ---
    purpose_words = get_words(descriptions.get('purpose_description', ''))
    best_info_match = (None, 0)
    for info_code, data in info_codes.items():
        info_desc_words = get_words(data['description'])
        score = len(purpose_words.intersection(info_desc_words))
        if score > best_info_match[1]:
            best_info_match = (info_code, score)

    if best_info_match[0] and best_info_match[1] > 0:
        found_codes["infoCode"] = best_info_match[0]
        
    logging.info(f"Derived codes from descriptions: {found_codes}")
    return found_codes

def format_dmc(parts):
    """Formats the final DMC string from its component parts."""
    sub_sys_code_formatted = parts.get("subSystemCode", DMC_DEFAULT_SUB_SYS_CODE).zfill(2)
    return (f'DMC-{DMC_MODEL_IDENT_CODE}-{DMC_SYSTEM_DIFF_CODE}-{parts.get("systemCode")}-'
            f'{sub_sys_code_formatted}-{DMC_ASSY_CODE}-'
            f'00A-{parts.get("infoCode")}A-{DMC_ITEM_LOCATION_CODE}')

def main():
    setup_environment()
    log_file_path = LOGS_DIR / f"run_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', handlers=[logging.FileHandler(log_file_path), logging.StreamHandler()])

    logging.info("--- Starting S1000D DMC Generation Process (v4) ---")
    
    sns_data = load_data_file(DATA_DIR / "sns.xml", parse_sns_from_file_object, "SNS XML")
    info_codes = load_data_file(DATA_DIR / "info_codes.json", json.load, "Info Codes JSON")

    if not sns_data or not info_codes: return

    files_to_process = [f for f in INPUT_DOCS_DIR.iterdir() if f.is_file() and f.suffix == '.docx' and not f.name.startswith('~')]
    
    if not files_to_process: logging.info("No new .docx files found.")
    
    log_summary = {"successful": [], "failed": []}
    
    for filepath in files_to_process:
        logging.info(f"--- Processing: {filepath.name} ---")
        full_text = extract_docx_structure(filepath)
        if not full_text:
            log_summary["failed"].append({"file": filepath.name, "reason": "Could not read document."})
            continue
        
        descriptions = get_descriptions_from_llm(full_text)
        if not descriptions:
            log_summary["failed"].append({"file": filepath.name, "reason": "LLM failed to analyze document."})
            continue

        final_codes = find_codes_from_descriptions(descriptions, sns_data, info_codes)
        final_dmc = format_dmc(final_codes)
        new_filename = f"{final_dmc}{filepath.suffix}"
        
        try:
            shutil.copy2(filepath, OUTPUT_DIR / new_filename)
            logging.info(f"Successfully created: {new_filename}")
            shutil.move(filepath, PROCESSED_DIR / filepath.name)
            log_summary["successful"].append({
                "original_file": filepath.name, "new_file": new_filename, "llm_analysis": descriptions, "derived_codes": final_codes
            })
        except Exception as e:
            logging.error(f"File operation failed for '{filepath.name}': {e}")
            log_summary["failed"].append({"file": filepath.name, "reason": f"File operation failed: {e}"})

    summary_path = LOGS_DIR / f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(summary_path, 'w') as f: json.dump(log_summary, f, indent=4)
    logging.info(f"--- Process Complete. Summary log saved to '{summary_path}' ---")

if __name__ == "__main__":
    main()